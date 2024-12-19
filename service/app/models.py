from docx import Document
from colorama import Fore
from time import sleep
from datetime import datetime
import os
import pandas as pd
from db_fns import set_db_column_value
import global_vars
#import streamlit as st

def processor(task):
    try:
        task_id = task['task_id']
        user_id = task['user_id']
        input_file = task['file_name']
        input_file_path = os.path.join(os.getcwd(),'data',f"{user_id}",'input',input_file)

        set_db_column_value(task_id, "task_status", "'processing'")
     
        set_db_column_value(task_id, "task_status_timestamp", f"'{datetime.now()}'")            
        set_db_column_value(task_id, "start_processing_timestamp", f"'{datetime.now()}'")   
        set_db_column_value(task_id, "task_progress", "0")      
        #global_vars.needs_to_reload = True       
        file_short_name = ".".join(input_file.split('.')[:-1])
        df = pd.read_excel(input_file_path, header=None)

        document = Document()
        document.add_paragraph(input_file)

        steps = len(df)
        for step, row in enumerate(df.itertuples()):
                sleep(1)
                document.add_paragraph(f"{int(row[1])}")
                set_db_column_value(task_id, 'task_progress', str(step/steps*100))
                set_db_column_value(task_id, 'task_progress_timestamp', f"'{datetime.now()}'")                
                print(f'{task_id} {int(step/steps*100)}')                
        # document.save(os.path.join(os.getcwd(), 'data', f'{user_id}', 'output', f'{file_short_name}.docx'))
        document.save(os.path.join(os.getcwd(), 'data', f'{user_id}', 'output', f'{input_file}.docx'))        

        set_db_column_value(task_id, 'task_progress', '100')  
        set_db_column_value(task_id, "task_status", "'ready'")  
        #global_vars.needs_to_reload = True
        print('*'*50, global_vars.needs_to_reload)
    except Exception as e:
        set_db_column_value(task_id, "task_status", "'error'")
        e=str(e).replace("'",'"')
        set_db_column_value(task_id, "task_rem", f"'{e}'")
    finally:
        set_db_column_value(task_id, 'task_progress_timestamp', f"'{datetime.now()}'")          
        print(Fore.GREEN, task, Fore.RESET)
       # global_vars.needs_to_reload = True
    #    st.rerun()
    #input('aaaa')





if __name__ == '__main__':
    #print(".".join("aaa.bbb.ccc".split(".")[:-1]))
    processor(2, "example_correct.xlsxx")
    processor(2, "example_with_err.xlsx")